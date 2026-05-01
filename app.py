from flask import Flask, send_file, request
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import io
import os

app = Flask(__name__)

def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_heading(doc, text, size=24, color='1C3D6E', bold=True, align='center'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def create_radar_chart(categories, values, title):
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    values_plot = values + values[:1]
    
    fig, ax = plt.subplots(1, 1, figsize=(6, 6), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('#1C3D6E')
    ax.set_facecolor('#1C3D6E')
    
    ax.plot(angles, values_plot, 'o-', linewidth=2, color='#FFD700')
    ax.fill(angles, values_plot, alpha=0.25, color='#FFD700')
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, color='white', size=11)
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(['20', '40', '60', '80', '100'], color='#FFD700', size=8)
    ax.grid(color='#FFD700', alpha=0.3)
    ax.spines['polar'].set_color('#FFD700')
    
    plt.title(title, color='#FFD700', size=14, fontweight='bold', pad=20)
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                facecolor='#1C3D6E')
    img_buffer.seek(0)
    plt.close()
    return img_buffer

def create_bar_chart(labels, values, title, color='#FFD700', target=None):
    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor('#1C3D6E')
    ax.set_facecolor('#0A1F3D')
    
    bars = ax.bar(labels, values, color=color, alpha=0.8, edgecolor='#FFD700')
    
    if target:
        ax.axhline(y=target, color='white', linestyle='--', linewidth=2, label=f'Target: {target}')
        ax.legend(facecolor='#1C3D6E', labelcolor='white')
    
    ax.set_title(title, color='#FFD700', fontsize=14, fontweight='bold')
    ax.tick_params(colors='white')
    ax.spines['bottom'].set_color('#FFD700')
    ax.spines['left'].set_color('#FFD700')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.5,
                str(val), ha='center', va='bottom', color='white', fontweight='bold')
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                facecolor='#1C3D6E')
    img_buffer.seek(0)
    plt.close()
    return img_buffer

@app.route('/generate', methods=['POST'])
def generate_report():
    data = request.json
    
    # Dati dal request
    peso = data.get('peso', '107.7')
    peso_obiettivo = data.get('peso_obiettivo', '85')
    peso_iniziale = data.get('peso_iniziale', '107.7')
    calorie_giorni = data.get('calorie_giorni', [1950, 1800, 2000, 1900, 2100, 1850, 1950])
    giorni_palestra = data.get('giorni_palestra', 3)
    acqua_media = data.get('acqua_media', 2.5)
    umore_medio = data.get('umore_medio', 75)
    stress_medio = data.get('stress_medio', 40)
    sonno_medio = data.get('sonno_medio', 70)
    streak = data.get('streak', 7)
    vita = data.get('vita', '-')
    collo = data.get('collo', '-')
    braccia = data.get('braccia', '-')
    petto = data.get('petto', '-')
    cosa_fatto_bene = data.get('cosa_fatto_bene', 'Hai seguito il piano pasti')
    cosa_migliorare = data.get('cosa_migliorare', 'Bevi più acqua')
    obiettivo = data.get('obiettivo', 'Palestra 3 volte questa settimana')
    messaggio_mister = data.get('messaggio_mister', 'Forza campione!')
    data_report = data.get('data', '01/05/2026')
    
    # Calcola rating
    rating_score = min(100, (giorni_palestra * 15) + (min(acqua_media/3.5, 1) * 20) + 
                       (umore_medio * 0.3) + (streak * 2))
    stelle = '⭐' * min(5, int(rating_score / 20) + 1)
    
    doc = Document()
    
    # ========== STILI GLOBALI ==========
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========== PAGINA 1 — COPERTINA ==========
    
    # Logo OC
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('⚽ OC ⚽')
    run.font.size = Pt(72)
    run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    run.font.bold = True
    
    add_heading(doc, 'OMAR COACH FC', size=36, color='1C3D6E')
    add_heading(doc, 'WEEKLY PERFORMANCE REPORT', size=20, color='FFD700')
    
    doc.add_paragraph()
    
    # Tabella info giocatore
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    cells_data = [
        ('👤 GIOCATORE', 'Omar Barhami'),
        ('👕 MAGLIA', f'#107 — Peso Iniziale'),
        ('🏆 STAGIONE', '2026'),
        ('📅 DATA REPORT', data_report),
    ]
    for i, (label, value) in enumerate(cells_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_background(row.cells[0], '1C3D6E')
        set_cell_background(row.cells[1], '0A1F3D')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
                    run.font.bold = True
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    add_heading(doc, f'RATING SETTIMANA: {stelle}', size=24, color='FFD700')
    add_heading(doc, f'Score: {int(rating_score)}/100', size=18, color='FFFFFF')
    
    doc.add_page_break()
    
    # ========== PAGINA 2 — STATISTICHE FISICHE ==========
    add_heading(doc, '⚽ STATISTICHE FISICHE', size=24, color='1C3D6E')
    doc.add_paragraph()
    
    # Calcola progresso
    try:
        kg_persi = float(peso_iniziale) - float(peso)
        progresso_pct = (kg_persi / (float(peso_iniziale) - float(peso_obiettivo))) * 100
        progresso_pct = max(0, min(100, progresso_pct))
    except:
        kg_persi = 0
        progresso_pct = 0
    
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    headers = ['STATISTICA', 'VALORE ATTUALE', 'OBIETTIVO']
    stats_data = [
        ('⚖️ Peso', f'{peso} kg', f'{peso_obiettivo} kg'),
        ('📉 Kg Persi', f'{kg_persi:.1f} kg', f'{float(peso_iniziale)-float(peso_obiettivo):.1f} kg'),
        ('📏 Giro Vita', f'{vita} cm', 'Migliorare'),
        ('📏 Giro Collo', f'{collo} cm', 'Migliorare'),
        ('📏 Giro Petto', f'{petto} cm', 'Migliorare'),
    ]
    
    # Header
    for j, h in enumerate(headers):
        cell = table2.rows[0].cells[j]
        cell.text = h
        set_cell_background(cell, '1C3D6E')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
                run.font.bold = True
    
    for i, (stat, val, obj) in enumerate(stats_data):
        row = table2.rows[i+1]
        row.cells[0].text = stat
        row.cells[1].text = val
        row.cells[2].text = obj
        set_cell_background(row.cells[0], '0A1F3D')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    doc.add_paragraph()
    add_heading(doc, f'🎯 Progresso verso obiettivo: {progresso_pct:.1f}%', 
                size=14, color='FFD700', align='left')
    
    # Barra progresso testuale
    barra_piena = int(progresso_pct / 5)
    barra = '█' * barra_piena + '░' * (20 - barra_piena)
    p = doc.add_paragraph()
    run = p.add_run(f'[{barra}] {progresso_pct:.1f}%')
    run.font.name = 'Courier New'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    
    doc.add_page_break()
    
    # ========== PAGINA 3 — PERFORMANCE ==========
    add_heading(doc, '📊 PERFORMANCE IN CAMPO', size=24, color='1C3D6E')
    
    # Grafico calorie
    giorni_labels = ['Lun', 'Mar', 'Mer', 'Gio', 'Ven', 'Sab', 'Dom']
    chart_cal = create_bar_chart(giorni_labels, calorie_giorni, 
                                  'Calorie Giornaliere vs Target (2000 kcal)', 
                                  target=2000)
    doc.add_picture(chart_cal, width=Inches(6))
    
    doc.add_paragraph()
    
    # Stats performance
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    perf_data = [
        ('⚽ Possesso Palla (Dieta)', f'{min(100, int((sum(1 for c in calorie_giorni if abs(c-2000)<200)/7)*100))}%'),
        ('🏃 Km Percorsi (Palestra)', f'{giorni_palestra}/7 giorni'),
        ('💧 Stamina (Acqua)', f'{acqua_media}L / 3.5L'),
    ]
    for i, (label, value) in enumerate(perf_data):
        row = table3.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_background(row.cells[0], '1C3D6E')
        set_cell_background(row.cells[1], '0A1F3D')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
                    run.font.bold = True
    
    doc.add_page_break()
    
    # ========== PAGINA 4 — ANALISI TATTICA ==========
    add_heading(doc, '🧠 ANALISI TATTICA', size=24, color='1C3D6E')
    
    # Grafico radar FIFA
    categories = ['Forza\nMentale', 'Disciplina\nDieta', 'Attività\nFisica', 
                  'Recupero\nSonno', 'Gestione\nStress', 'Idratazione']
    values = [
        umore_medio,
        min(100, int((sum(1 for c in calorie_giorni if abs(c-2000)<300)/7)*100)),
        min(100, giorni_palestra * 33),
        sonno_medio,
        max(0, 100 - stress_medio),
        min(100, int((acqua_media/3.5)*100))
    ]
    
    radar_chart = create_radar_chart(categories, values, 'ATTRIBUTI FIFA - Omar Barhami')
    doc.add_picture(radar_chart, width=Inches(5))
    
    doc.add_page_break()
    
    # ========== PAGINA 5 — SCOUT REPORT ==========
    add_heading(doc, '🔍 SCOUT REPORT', size=24, color='1C3D6E')
    doc.add_paragraph()
    
    add_heading(doc, '✅ COSA HAI FATTO BENE', size=14, color='00B050', align='left')
    p = doc.add_paragraph()
    run = p.add_run(cosa_fatto_bene)
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    add_heading(doc, '⚠️ AREE DI MIGLIORAMENTO', size=14, color='FF6600', align='left')
    p = doc.add_paragraph()
    run = p.add_run(cosa_migliorare)
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== PAGINA 6 — PROSSIMA PARTITA ==========
    add_heading(doc, '🎯 PROSSIMA PARTITA', size=24, color='1C3D6E')
    doc.add_paragraph()
    
    add_heading(doc, '📋 OBIETTIVO SETTIMANA', size=16, color='1C3D6E', align='left')
    p = doc.add_paragraph()
    run = p.add_run(obiettivo)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    
    doc.add_paragraph()
    add_heading(doc, '💬 IL MISTER DICE:', size=16, color='1C3D6E', align='left')
    p = doc.add_paragraph()
    run = p.add_run(f'"{messaggio_mister}"')
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1C, 0x3D, 0x6E)
    
    doc.add_page_break()
    
    # ========== PAGINA 7 — HALL OF FAME ==========
    add_heading(doc, '🏅 HALL OF FAME', size=24, color='1C3D6E')
    doc.add_paragraph()
    
    add_heading(doc, f'🔥 STREAK ATTUALE: {streak} GIORNI', size=18, color='FFD700')
    doc.add_paragraph()
    
    # Trofei
    trofei = []
    if giorni_palestra >= 3:
        trofei.append('🏆 HAT-TRICK — 3+ giorni in palestra questa settimana!')
    if streak >= 7:
        trofei.append('🥇 CAPITANO — 7 giorni consecutivi streak!')
    if kg_persi > 0:
        trofei.append(f'⚽ GOL DEL MESE — Persi {kg_persi:.1f} kg totali!')
    if acqua_media >= 3:
        trofei.append('💧 IDRATAZIONE PRO — Media acqua superiore a 3L!')
    
    if trofei:
        add_heading(doc, '🏆 TROFEI SBLOCCATI QUESTA SETTIMANA:', size=16, color='1C3D6E', align='left')
        for trofeo in trofei:
            p = doc.add_paragraph()
            run = p.add_run(trofeo)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    else:
        add_heading(doc, '💪 Continua così — i trofei arrivano!', size=16, color='FFD700', align='left')
    
    # Salva in memoria
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return send_file(
        doc_buffer,
        as_attachment=True,
        download_name='Omar_Coach_FC_Report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
